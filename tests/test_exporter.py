"""
tests/test_exporter.py — Unit tests for export_engine.exporter.

Covers:
  - Timestamp formatting helpers (SRT and VTT)
  - SRT export: file created, correct structure
  - VTT export: file created, starts with WEBVTT header
  - PDF export: valid output, tier highlighting, empty transcript, output_dir isolation
  - DOCX export: valid output, tier highlighting, empty transcript, output_dir isolation
  - ZIP export: honours output_dir for sidecars and does not contaminate global dir
"""

from __future__ import annotations

import io
import re
import zipfile

from export_engine.exporter import (
    _md_to_html,
    _seconds_to_srt_ts,
    _seconds_to_vtt_ts,
    export_docx,
    export_pdf,
    export_srt,
    export_vtt,
    export_zip,
)

# ─────────────────────────────────────────────────────────────────────────────
# Timestamp helpers
# ─────────────────────────────────────────────────────────────────────────────


class TestTimestampFormatting:
    def test_srt_zero_seconds(self):
        assert _seconds_to_srt_ts(0.0) == "00:00:00,000"

    def test_srt_one_hour_one_minute_one_second_half(self):
        assert _seconds_to_srt_ts(3661.5) == "01:01:01,500"

    def test_srt_sub_second_precision(self):
        assert _seconds_to_srt_ts(0.123) == "00:00:00,123"

    def test_vtt_zero_seconds(self):
        assert _seconds_to_vtt_ts(0.0) == "00:00:00.000"

    def test_vtt_uses_period_not_comma(self):
        """VTT format uses '.' as millisecond separator, not ','."""
        ts = _seconds_to_vtt_ts(3661.5)
        assert "." in ts
        assert "," not in ts

    def test_srt_uses_comma_not_period(self):
        ts = _seconds_to_srt_ts(3661.5)
        assert "," in ts
        assert ts.count(".") == 0


# ─────────────────────────────────────────────────────────────────────────────
# SRT export
# ─────────────────────────────────────────────────────────────────────────────


def _mock_whisper_result():
    return {
        "text": "Hello world. This is a test.",
        "segments": [
            {"start": 0.0, "end": 2.5, "text": " Hello world."},
            {"start": 2.5, "end": 5.0, "text": " This is a test."},
        ],
        "language": "en",
    }


class TestSRTExport:
    def test_srt_file_is_created(self):
        result = export_srt(_mock_whisper_result(), stem="test_srt")
        assert result.exists()

    def test_srt_file_has_correct_extension(self):
        result = export_srt(_mock_whisper_result(), stem="test_srt_ext")
        assert result.suffix == ".srt"

    def test_srt_has_valid_cue_structure(self):
        result = export_srt(_mock_whisper_result(), stem="test_srt_seq")
        content = result.read_text(encoding="utf-8").strip()

        cues = [block.splitlines() for block in content.split("\n\n") if block.strip()]
        assert len(cues) == 2

        ts_pattern = re.compile(
            r"^\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3}$"
        )
        for idx, cue in enumerate(cues, start=1):
            assert cue[0] == str(idx)
            assert ts_pattern.match(cue[1])
            assert cue[2].strip()

    def test_srt_empty_segments_produces_empty_file(self):
        result = export_srt({"segments": []}, stem="test_srt_empty")
        content = result.read_text(encoding="utf-8")
        assert content.strip() == ""


# ─────────────────────────────────────────────────────────────────────────────
# VTT export
# ─────────────────────────────────────────────────────────────────────────────


class TestVTTExport:
    def test_vtt_file_is_created(self):
        result = export_vtt(_mock_whisper_result(), stem="test_vtt")
        assert result.exists()

    def test_vtt_starts_with_webvtt_header(self):
        result = export_vtt(_mock_whisper_result(), stem="test_vtt_header")
        content = result.read_text(encoding="utf-8")
        assert content.startswith("WEBVTT")

    def test_vtt_has_valid_header_and_cue_structure(self):
        result = export_vtt(_mock_whisper_result(), stem="test_vtt_period")
        content = result.read_text(encoding="utf-8")

        lines = content.splitlines()
        assert lines[0] == "WEBVTT"

        ts_pattern = re.compile(
            r"^\d{2}:\d{2}:\d{2}\.\d{3} --> \d{2}:\d{2}:\d{2}\.\d{3}$"
        )
        ts_lines = [ln for ln in lines if "-->" in ln]
        assert len(ts_lines) == 2
        for line in ts_lines:
            assert ts_pattern.match(line)
            assert "," not in line

        assert "\n\n" in content, "VTT cues should be separated by blank lines"

    def test_vtt_contains_transcript_text(self):
        result = export_vtt(_mock_whisper_result(), stem="test_vtt_text")
        content = result.read_text(encoding="utf-8")
        assert "Hello world" in content or "hello world" in content.lower()


# ─────────────────────────────────────────────────────────────────────────────
# JSON transcript bundle
# ─────────────────────────────────────────────────────────────────────────────


class TestExportTranscriptBundle:
    def _make_votes(self):
        from consensus_merger.alignment import WordVote

        return [
            WordVote(
                word="hello",
                count=4,
                total=4,
                confidence=1.0,
                tier="HIGH",
                variants=["hello"],
            ),
            WordVote(
                word="world",
                count=2,
                total=4,
                confidence=0.5,
                tier="MEDIUM",
                variants=["world", "word"],
            ),
            WordVote(
                word="garbl",
                count=1,
                total=4,
                confidence=0.25,
                tier="LOW",
                variants=["garbl"],
            ),
        ]

    def _make_transcripts(self):
        return {
            "original": {
                "text": "hello world garbl",
                "language": "en",
                "model": "base",
                "device": "cpu",
            },
            "highpass": {
                "text": "hello word garbl",
                "language": "en",
                "model": "base",
                "device": "cpu",
            },
        }

    def test_bundle_file_created(self, tmp_path):
        from export_engine.exporter import export_transcript_bundle

        path = export_transcript_bundle(
            self._make_transcripts(), self._make_votes(), "test", output_dir=tmp_path
        )
        assert path.exists()
        assert path.name == "test_bundle.json"

    def test_bundle_structure(self, tmp_path):
        import json

        from export_engine.exporter import export_transcript_bundle

        path = export_transcript_bundle(
            self._make_transcripts(), self._make_votes(), "test", output_dir=tmp_path
        )
        data = json.loads(path.read_text(encoding="utf-8"))
        assert "meta" in data
        assert data["meta"]["stem"] == "test"
        assert "variants" in data
        assert "original" in data["variants"]
        assert data["variants"]["original"]["text"] == "hello world garbl"
        assert "consensus" in data
        assert len(data["consensus"]) == 3
        assert data["consensus"][0] == {
            "word": "hello",
            "tier": "HIGH",
            "confidence": 1.0,
            "count": 4,
            "total": 4,
            "variants": ["hello"],
        }
        assert "statistics" in data
        assert data["statistics"]["high"] == 1
        assert data["statistics"]["medium"] == 1
        assert data["statistics"]["low"] == 1
        assert data["statistics"]["total_words"] == 3

    def test_bundle_meta_versioned(self, tmp_path):
        """meta carries the producing Chorus version and the schema revision."""
        import json
        from pathlib import Path

        from export_engine.exporter import (
            BUNDLE_SCHEMA_VERSION,
            export_transcript_bundle,
        )

        path = export_transcript_bundle(
            self._make_transcripts(), self._make_votes(), "test", output_dir=tmp_path
        )
        data = json.loads(path.read_text(encoding="utf-8"))
        repo_root = Path(__file__).resolve().parent.parent
        expected_version = (repo_root / "VERSION").read_text(encoding="utf-8").strip()
        assert data["meta"]["chorus_version"] == expected_version
        assert data["meta"]["schema_version"] == BUNDLE_SCHEMA_VERSION
        assert data["meta"]["schema_version"] == 1

    def test_bundle_matches_documented_contract(self, tmp_path):
        """docs/CHORUS_FOR_LLMS.md §5 and the real bundle must agree.

        If a bundle field is added, renamed, or removed without updating the
        documented schema — or vice versa — this test fails, pointing at both
        export_engine/exporter.py::export_transcript_bundle and
        docs/CHORUS_FOR_LLMS.md §5.
        """
        import json
        import re
        from pathlib import Path

        from export_engine.exporter import export_transcript_bundle

        doc = (
            Path(__file__).resolve().parent.parent / "docs" / "CHORUS_FOR_LLMS.md"
        ).read_text(encoding="utf-8")
        section = doc.split("## 5.")[1].split("\n## ")[0]
        fence = re.search(r"```json\n(.*?)```", section, re.S)
        assert fence, "no JSON example found in CHORUS_FOR_LLMS.md §5"
        example = fence.group(1)

        path = export_transcript_bundle(
            self._make_transcripts(), self._make_votes(), "test", output_dir=tmp_path
        )
        data = json.loads(path.read_text(encoding="utf-8"))

        # Every key the real bundle produces must appear in the documented
        # example (as a quoted string), at each structural level.
        for key in data:
            assert f'"{key}"' in example, f"bundle key {key!r} missing from doc §5"
        for key in data["meta"]:
            assert f'"{key}"' in example, f"meta key {key!r} missing from doc §5"
        for key in data["consensus"][0]:
            assert f'"{key}"' in example, f"consensus key {key!r} missing from doc §5"
        for key in data["statistics"]:
            assert f'"{key}"' in example, f"statistics key {key!r} missing from doc §5"

        # And the documented consensus-entry keys must all exist in the real
        # bundle, so removals/renames are caught in both directions.
        documented_consensus_keys = {
            "word",
            "tier",
            "confidence",
            "count",
            "total",
            "variants",
        }
        assert set(data["consensus"][0]) == documented_consensus_keys
        documented_meta_keys = {
            "stem",
            "source_filename",
            "generated_at",
            "chorus_version",
            "schema_version",
        }
        assert set(data["meta"]) == documented_meta_keys

    def test_bundle_in_pipeline_output(self, tmp_path):
        """run_pipeline should return bundle_path in its result dict."""
        from unittest.mock import patch

        from tests.test_integration import (
            _generate_sine_wav,
            _mock_run_transcription_pass,
        )

        audio = _generate_sine_wav(tmp_path / "audio.wav")
        with patch(
            "pipeline_runner.run_transcription_pass",
            side_effect=_mock_run_transcription_pass,
        ):
            from pipeline_runner import run_pipeline

            result = run_pipeline(
                audio_path=audio, language="en", output_dir=tmp_path / "out"
            )
        assert "bundle_path" in result
        assert result["bundle_path"].exists()
        assert result["bundle_path"].name.endswith("_bundle.json")


# ─────────────────────────────────────────────────────────────────────────────
# Best-guess transcript export
# ─────────────────────────────────────────────────────────────────────────────


class TestExportBestGuess:
    def _make_mixed_tier_votes(self):
        from consensus_merger.alignment import WordVote

        return [
            WordVote(
                word="hello",
                count=4,
                total=4,
                confidence=1.0,
                tier="HIGH",
                variants=["hello"],
            ),
            WordVote(
                word="world",
                count=2,
                total=4,
                confidence=0.5,
                tier="MEDIUM",
                variants=["world", "word"],
            ),
            WordVote(
                word="today",
                count=1,
                total=4,
                confidence=0.25,
                tier="LOW",
                variants=["today"],
            ),
        ]

    def _render(self, votes, tmp_path):
        from consensus_merger.renderer import render_consensus

        transcripts_meta = {
            "original": {"text": "hello world today", "model": "base", "language": "en"}
        }
        return render_consensus(votes, "test", transcripts_meta, consensus_dir=tmp_path)

    def test_best_guess_contains_high_agreement_word_no_markup(self, tmp_path):
        """Best-guess file must contain the winning word at every tier, with
        no brackets, confidence annotations, or statistics lines."""
        from export_engine.exporter import export_best_guess

        votes = self._make_mixed_tier_votes()
        consensus_path = self._render(votes, tmp_path)

        out_path = export_best_guess(consensus_path, "test", output_dir=tmp_path)
        content = out_path.read_text(encoding="utf-8")

        assert out_path.name == "test_best_guess.txt"
        # The winning word at every position (including MEDIUM/LOW) is present.
        assert "hello" in content
        assert "world" in content
        assert "today" in content
        # No confidence markup of any kind.
        assert "[" not in content
        assert "?]" not in content
        assert "==" not in content
        assert "~~" not in content
        # No statistics/legend lines leaked into the transcript.
        assert "HIGH" not in content
        assert "LOW" not in content
        assert "%" not in content

    def test_best_guess_empty_transcript_produces_empty_file(self, tmp_path):
        """Silence (no votes) must produce an empty file, not raise."""
        from export_engine.exporter import export_best_guess

        consensus_path = self._render([], tmp_path)
        out_path = export_best_guess(consensus_path, "silent", output_dir=tmp_path)

        assert out_path.exists()
        assert out_path.read_text(encoding="utf-8") == ""

    def test_best_guess_honours_output_dir(self, tmp_path):
        """File must land under the supplied output_dir, not the global dir."""
        from export_engine.exporter import export_best_guess

        votes = self._make_mixed_tier_votes()
        isolated_dir = tmp_path / "isolated"
        isolated_dir.mkdir()
        consensus_path = self._render(votes, isolated_dir)

        out_path = export_best_guess(consensus_path, "test", output_dir=isolated_dir)

        assert out_path.parent == isolated_dir


# ─────────────────────────────────────────────────────────────────────────────
# Markdown → HTML conversion helper
# ─────────────────────────────────────────────────────────────────────────────


class TestMdToHtml:
    def test_converts_heading_and_maps_medium_highlight_to_mark_tag(self):
        html = _md_to_html("# Title\n\nhello ==world== plain")
        assert "<h1>Title</h1>" in html
        assert "<mark>world</mark>" in html

    def test_empty_markdown_does_not_raise(self):
        html = _md_to_html("")
        assert isinstance(html, str)

    def test_converts_strikethrough_tilde_syntax_to_del_tag(self):
        """Strikethrough `~~word~~` must convert to `<del>word</del>`."""
        html = _md_to_html("hello ~~garbl~~ world")
        assert "<del>garbl</del>" in html
        # Literal tildes must NOT appear in the output
        assert "~~" not in html

    def test_strikethrough_regression_medium_highlight_still_works(self):
        """Regression guard: ==word== must still map to <mark> after strikethrough fix."""
        html = _md_to_html("hello ==medium== word")
        assert "<mark>medium</mark>" in html

    def test_single_tilde_passes_through_unchanged(self):
        """A single tilde (not doubled) must not be affected by strikethrough processing."""
        html = _md_to_html("hello ~ world")
        # Should not crash or mangle; the tilde survives (possibly escaped)
        assert "hello" in html
        assert "world" in html


# ─────────────────────────────────────────────────────────────────────────────
# Per-run consensus threshold reflected in the rendered legend
# ─────────────────────────────────────────────────────────────────────────────


class TestRenderConsensusThresholdLegend:
    def _votes(self):
        from consensus_merger.alignment import WordVote

        return [
            WordVote(
                word="hello",
                count=4,
                total=4,
                confidence=1.0,
                tier="HIGH",
                variants=["hello"],
            ),
        ]

    def _render(self, tmp_path, **kwargs):
        from consensus_merger.renderer import render_consensus

        transcripts_meta = {
            "original": {"text": "hello", "model": "base", "language": "en"}
        }
        return render_consensus(
            self._votes(), "test", transcripts_meta, consensus_dir=tmp_path, **kwargs
        )

    def test_custom_threshold_reflected_in_legend(self, tmp_path):
        """Passing consensus_threshold=0.9 must render '90' in the legend and
        must not render the default '75 %' bar."""
        consensus_path = self._render(tmp_path, consensus_threshold=0.9)
        text = consensus_path.read_text(encoding="utf-8")
        assert "≥ 90 %" in text
        assert "75 %" not in text

    def test_default_threshold_omitted_matches_current_wording(self, tmp_path):
        """Omitting consensus_threshold must reproduce the existing default
        wording exactly (regression guard)."""
        consensus_path = self._render(tmp_path)
        text = consensus_path.read_text(encoding="utf-8")
        assert "≥ 75 %" in text


# ─────────────────────────────────────────────────────────────────────────────
# PDF export
# ─────────────────────────────────────────────────────────────────────────────


def _mixed_tier_votes():
    from consensus_merger.alignment import WordVote

    return [
        WordVote(
            word="hello",
            count=4,
            total=4,
            confidence=1.0,
            tier="HIGH",
            variants=["hello"],
        ),
        WordVote(
            word="world",
            count=2,
            total=4,
            confidence=0.5,
            tier="MEDIUM",
            variants=["world", "word"],
        ),
        WordVote(
            word="garbl",
            count=1,
            total=4,
            confidence=0.25,
            tier="LOW",
            variants=["garbl"],
        ),
    ]


def _render_mixed_tier_consensus(tmp_path, votes=None):
    from consensus_merger.renderer import render_consensus

    votes = _mixed_tier_votes() if votes is None else votes
    transcripts_meta = {
        "original": {"text": "hello world garbl", "model": "base", "language": "en"}
    }
    return render_consensus(votes, "test", transcripts_meta, consensus_dir=tmp_path)


class TestPDFExport:
    def test_pdf_export_creates_valid_pdf(self, tmp_path):
        consensus_path = _render_mixed_tier_consensus(tmp_path)
        out_path = export_pdf(consensus_path, "test", output_dir=tmp_path)

        assert out_path.exists()
        assert out_path.suffix == ".pdf"
        data = out_path.read_bytes()
        assert data.startswith(b"%PDF-")
        assert len(data) > 500  # non-trivial size, not a truncated/empty file

    def test_pdf_export_preserves_tier_highlighting(self, tmp_path, monkeypatch):
        """MEDIUM markup must reach the HTML handed to WeasyPrint as a <mark>
        tag, and LOW-tier text must not be silently dropped.

        Note: the LOW-tier word is *not* wrapped in a <del> tag here — see the
        bug flagged in the PR description — but it must still be present.
        """
        captured: dict[str, str] = {}

        import weasyprint  # type: ignore

        class _SpyHTML(weasyprint.HTML):
            def __init__(self, *, string, **kwargs):
                captured["html"] = string
                super().__init__(string=string, **kwargs)

        # export_pdf imports HTML from weasyprint lazily inside the function
        # body, so the module attribute must be patched rather than the
        # (non-existent) name in export_engine.exporter's namespace.
        monkeypatch.setattr(weasyprint, "HTML", _SpyHTML)

        consensus_path = _render_mixed_tier_consensus(tmp_path)
        out_path = export_pdf(consensus_path, "test", output_dir=tmp_path)

        assert out_path.read_bytes().startswith(b"%PDF-")
        assert "<mark>world</mark>" in captured["html"]
        assert "garbl" in captured["html"]

    def test_pdf_export_low_tier_strikethrough_converted_to_del(
        self, tmp_path, monkeypatch
    ):
        """LOW-confidence word in Markdown as **~~word~~** must be converted to
        <del>word</del> in HTML so the CSS rule `del strong, strong del` can style it.
        Literal `~~` must not appear in the HTML.
        """
        captured: dict[str, str] = {}

        import weasyprint  # type: ignore

        class _SpyHTML(weasyprint.HTML):
            def __init__(self, *, string, **kwargs):
                captured["html"] = string
                super().__init__(string=string, **kwargs)

        monkeypatch.setattr(weasyprint, "HTML", _SpyHTML)

        consensus_path = _render_mixed_tier_consensus(tmp_path)
        out_path = export_pdf(consensus_path, "test", output_dir=tmp_path)

        assert out_path.read_bytes().startswith(b"%PDF-")
        # LOW-confidence word must be wrapped in <del> tag
        assert "<del>garbl</del>" in captured["html"]
        # Literal tilde characters must not appear in the HTML
        assert "~~" not in captured["html"]

    def test_pdf_export_empty_transcript_does_not_crash(self, tmp_path):
        """Silence (no votes) must still produce a valid PDF, not raise."""
        consensus_path = _render_mixed_tier_consensus(tmp_path, votes=[])
        out_path = export_pdf(consensus_path, "silent", output_dir=tmp_path)

        assert out_path.exists()
        assert out_path.read_bytes().startswith(b"%PDF-")

    def test_pdf_export_honours_output_dir(self, tmp_path, monkeypatch):
        """File must land under the supplied output_dir, not the global dir."""
        global_dir = tmp_path / "global"
        global_dir.mkdir()
        monkeypatch.setattr("export_engine.exporter.CONSENSUS_DIR", global_dir)

        isolated_dir = tmp_path / "isolated"
        isolated_dir.mkdir()
        consensus_path = _render_mixed_tier_consensus(isolated_dir)

        out_path = export_pdf(consensus_path, "test", output_dir=isolated_dir)

        assert out_path.parent == isolated_dir
        assert not any(global_dir.iterdir())


# ─────────────────────────────────────────────────────────────────────────────
# DOCX export
# ─────────────────────────────────────────────────────────────────────────────


class TestDOCXExport:
    def test_docx_export_creates_valid_document_with_expected_structure(self, tmp_path):
        from docx import Document

        consensus_path = _render_mixed_tier_consensus(tmp_path)
        out_path = export_docx(consensus_path, "test", output_dir=tmp_path)

        assert out_path.exists()
        assert out_path.suffix == ".docx"

        doc = Document(str(out_path))
        headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert "Chorus — Consensus Transcript" in headings
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "hello" in body_text
        assert "world" in body_text
        assert "garbl" in body_text

    def test_docx_export_preserves_tier_highlighting(self, tmp_path):
        """HIGH/MEDIUM/LOW markup must round-trip into styled runs: MEDIUM as a
        yellow highlight, LOW as a red highlight with strikethrough."""
        from docx import Document
        from docx.enum.text import WD_COLOR_INDEX

        consensus_path = _render_mixed_tier_consensus(tmp_path)
        out_path = export_docx(consensus_path, "test", output_dir=tmp_path)

        doc = Document(str(out_path))
        all_runs = [r for p in doc.paragraphs for r in p.runs]

        high_runs = [r for r in all_runs if "hello" in r.text]
        assert high_runs
        assert high_runs[0].font.highlight_color is None

        medium_runs = [r for r in all_runs if "world" in r.text]
        assert medium_runs
        assert medium_runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW

        low_runs = [r for r in all_runs if "garbl" in r.text]
        assert low_runs
        assert low_runs[0].font.highlight_color == WD_COLOR_INDEX.RED
        assert low_runs[0].font.strike is True

    def test_docx_export_empty_transcript_does_not_crash(self, tmp_path):
        """Silence (no votes) must still produce a valid, openable document."""
        from docx import Document

        consensus_path = _render_mixed_tier_consensus(tmp_path, votes=[])
        out_path = export_docx(consensus_path, "silent", output_dir=tmp_path)

        assert out_path.exists()
        doc = Document(str(out_path))  # must open without raising
        assert len(doc.paragraphs) > 0

    def test_docx_export_honours_output_dir(self, tmp_path, monkeypatch):
        """File must land under the supplied output_dir, not the global dir."""
        global_dir = tmp_path / "global"
        global_dir.mkdir()
        monkeypatch.setattr("export_engine.exporter.CONSENSUS_DIR", global_dir)

        isolated_dir = tmp_path / "isolated"
        isolated_dir.mkdir()
        consensus_path = _render_mixed_tier_consensus(isolated_dir)

        out_path = export_docx(consensus_path, "test", output_dir=isolated_dir)

        assert out_path.parent == isolated_dir
        assert not any(global_dir.iterdir())


# ─────────────────────────────────────────────────────────────────────────────
# ZIP export with output_dir isolation
# ─────────────────────────────────────────────────────────────────────────────


class TestZipExportOutputDirIsolation:
    """Test that export_zip respects output_dir and does not leak to global dir."""

    def _make_mock_whisper_result(self) -> dict:
        """Minimal Whisper result for testing."""
        return {
            "text": "test recording",
            "segments": [{"start": 0.0, "end": 1.0, "text": " test recording"}],
            "language": "en",
        }

    def test_zip_reads_sidecars_from_output_dir(self, tmp_path):
        """ZIP should include sidecars written to the specified output_dir."""
        # Create a consensus markdown in output_dir
        output_dir = tmp_path / "outputs"
        output_dir.mkdir()
        consensus_path = output_dir / "test_consensus.md"
        consensus_path.write_text("# Test Consensus\n\nHello world.\n")

        # Create the sidecars in output_dir
        speakers_path = output_dir / "test_speakers.json"
        speakers_path.write_text('{"SPEAKER_00": "Alice"}')

        ai_context_path = output_dir / "test_ai_context.md"
        ai_context_path.write_text("# AI Context\n\nGood recording.\n")

        diarised_path = output_dir / "test_diarised.md"
        diarised_path.write_text("# Diarised\n\nAlice: Hello world.\n")

        # Build ZIP
        zip_bytes = export_zip(
            consensus_md_path=consensus_path,
            whisper_result=self._make_mock_whisper_result(),
            stem="test",
            output_dir=output_dir,
        )

        # Verify ZIP contains all sidecars
        zf = zipfile.ZipFile(io.BytesIO(zip_bytes))
        names = zf.namelist()
        assert any(
            "_speakers.json" in n for n in names
        ), "Missing speakers sidecar in ZIP"
        assert any("_ai_context.md" in n for n in names), "Missing AI context in ZIP"
        assert any(
            "_diarised.md" in n for n in names
        ), "Missing diarised sidecar in ZIP"

    def test_zip_does_not_pick_up_stale_sidecars_from_global_dir(
        self, tmp_path, monkeypatch
    ):
        """ZIP from isolated output_dir should not include stale files from global dir."""
        # Mock global CONSENSUS_DIR to tmp_path / "global"
        global_dir = tmp_path / "global"
        global_dir.mkdir()
        monkeypatch.setattr("export_engine.exporter.CONSENSUS_DIR", global_dir)

        # Create stale files in global dir
        stale_speakers = global_dir / "test_speakers.json"
        stale_speakers.write_text('{"SPEAKER_00": "OldName"}')

        stale_ai_context = global_dir / "test_ai_context.md"
        stale_ai_context.write_text("# Old AI Context\n\nStale data.\n")

        # Create fresh consensus in isolated output_dir
        isolated_dir = tmp_path / "isolated"
        isolated_dir.mkdir()
        consensus_path = isolated_dir / "test_consensus.md"
        consensus_path.write_text("# Test Consensus\n\nFresh recording.\n")

        # Build ZIP from isolated_dir (no sidecars there)
        zip_bytes = export_zip(
            consensus_md_path=consensus_path,
            whisper_result=self._make_mock_whisper_result(),
            stem="test",
            output_dir=isolated_dir,
        )

        # Verify ZIP does NOT contain the stale sidecars from global dir
        zf = zipfile.ZipFile(io.BytesIO(zip_bytes))
        names = zf.namelist()
        assert not any(
            "_speakers.json" in n for n in names
        ), "ZIP should not include stale speakers sidecar from global dir"
        assert not any(
            "_ai_context.md" in n for n in names
        ), "ZIP should not include stale AI context from global dir"

    def test_zip_contains_consensus_and_plaintext(self, tmp_path):
        """ZIP should always include consensus markdown and plaintext variants."""
        output_dir = tmp_path / "outputs"
        output_dir.mkdir()
        consensus_path = output_dir / "test_consensus.md"
        consensus_path.write_text(
            "# Consensus\n\n## Consensus Transcript\n\n**~~guess~~**[^1]\n\n"
        )

        zip_bytes = export_zip(
            consensus_md_path=consensus_path,
            whisper_result=self._make_mock_whisper_result(),
            stem="test",
            output_dir=output_dir,
        )

        zf = zipfile.ZipFile(io.BytesIO(zip_bytes))
        names = zf.namelist()
        # Should have consensus markdown
        assert any("consensus.md" in n for n in names)
        # Should have plaintext variants
        assert any("most_likely" in n for n in names)
