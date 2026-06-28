"""
export_engine/exporter.py — Multi-format export module.

Converts the Chorus consensus Markdown document and/or the Whisper segment
data into the following output formats:

  - PDF   : via WeasyPrint (HTML → PDF pipeline, preserves highlighting)
  - DOCX  : via python-docx (native Word document with styled runs)
  - SRT   : SubRip subtitle format (timed captions, widely supported)
  - VTT   : WebVTT subtitle format (HTML5 <track> compatible)

All exports are written to outputs/consensus/ alongside the .md source.

Usage
─────
    from export_engine.exporter import export_all

    paths = export_all(
        consensus_md_path=Path("outputs/consensus/recording_consensus.md"),
        whisper_result=transcripts["original"],   # for SRT/VTT timestamps
        stem="recording",
        formats=["pdf", "docx", "srt", "vtt"],
    )
"""

from __future__ import annotations

import io
import logging
import re
import zipfile
from datetime import timedelta
from pathlib import Path
from typing import Any

from config import CONSENSUS_DIR

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────


def _seconds_to_srt_ts(seconds: float) -> str:
    """Format seconds as SRT timestamp: HH:MM:SS,mmm"""
    td = timedelta(seconds=seconds)
    total_s = int(td.total_seconds())
    ms = int((td.total_seconds() - total_s) * 1000)
    h, rem = divmod(total_s, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _seconds_to_vtt_ts(seconds: float) -> str:
    """Format seconds as WebVTT timestamp: HH:MM:SS.mmm"""
    return _seconds_to_srt_ts(seconds).replace(",", ".")


def _strip_md_markup(text: str) -> str:
    """Remove Markdown decorators for plain-text subtitle output."""
    text = re.sub(r"==(.+?)==", r"\1", text)  # ==highlight==
    text = re.sub(r"\*\*~~(.+?)~~\*\*\[.*?\]", r"\1", text)  # **~~word~~**[...]
    text = re.sub(r"[*_~`#>]", "", text)
    return text.strip()


def _md_to_html(md_text: str) -> str:
    """Convert Markdown to HTML with basic highlight support."""
    try:
        import markdown as md_lib

        html = md_lib.markdown(md_text, extensions=["tables", "fenced_code"])
    except ImportError:
        # Minimal fallback
        html = "<pre>" + md_text.replace("&", "&amp;").replace("<", "&lt;") + "</pre>"

    # Map ==highlight== → <mark> (not handled by python-markdown by default)
    html = re.sub(r"==(.+?)==", r"<mark>\1</mark>", html)
    return html


# ─────────────────────────────────────────────────────────────────────────────
# PDF export
# ─────────────────────────────────────────────────────────────────────────────


def export_pdf(
    consensus_md_path: Path,
    stem: str,
    output_dir: Path | None = None,
    source_filename: str | None = None,
) -> Path:
    """
    Export the consensus Markdown document to PDF via WeasyPrint.

    The Markdown is first converted to HTML, styled with an embedded
    CSS sheet that preserves the confidence highlighting colours, and
    then rendered to PDF by WeasyPrint.

    Parameters
    ----------
    consensus_md_path : Path
        Path to the consensus ``.md`` file.
    stem : str
        Base filename stem for the output file.

    Returns
    -------
    Path
        Path to the written ``.pdf`` file.
    """
    try:
        from weasyprint import CSS, HTML  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "WeasyPrint is required for PDF export. "
            "Install with: pip install weasyprint"
        ) from exc

    md_text = consensus_md_path.read_text(encoding="utf-8")
    html_body = _md_to_html(md_text)

    css = CSS(string="""
        @page { margin: 2cm; }
        body {
            font-family: 'Georgia', serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #1a1a2e;
        }
        h1 { font-size: 20pt; border-bottom: 2px solid #0f3460; padding-bottom: 6px; }
        h2 { font-size: 14pt; color: #0f3460; margin-top: 1.5em; }
        h3 { font-size: 12pt; }
        mark {
            background-color: #fff3cd;
            color: #856404;
            padding: 1px 3px;
            border-radius: 3px;
        }
        del strong, strong del {
            background-color: #f8d7da;
            color: #721c24;
            text-decoration: line-through;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
            font-size: 10pt;
        }
        th { background: #0f3460; color: white; padding: 6px 10px; }
        td { border: 1px solid #dee2e6; padding: 5px 10px; }
        blockquote {
            border-left: 4px solid #0f3460;
            margin: 0.5em 0;
            padding: 0.5em 1em;
            background: #f0f2f6;
            color: #444;
        }
        code {
            background: #f0f2f6;
            padding: 1px 4px;
            border-radius: 3px;
            font-size: 10pt;
        }
        hr { border: none; border-top: 1px solid #dee2e6; margin: 1.5em 0; }
    """)

    full_html = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="utf-8"><title>Chorus Consensus Transcript</title></head>
<body>{html_body}</body>
</html>"""

    target_dir = output_dir or CONSENSUS_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / f"{stem}_consensus.pdf"
    HTML(string=full_html).write_pdf(str(out_path), stylesheets=[css])
    logger.info("PDF export written → %s", out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# DOCX export
# ─────────────────────────────────────────────────────────────────────────────


def export_docx(
    consensus_md_path: Path,
    stem: str,
    output_dir: Path | None = None,
    source_filename: str | None = None,
) -> Path:
    """
    Export the consensus Markdown document to a styled DOCX file.

    Confidence tiers are mapped to Word character styles:
      HIGH   → plain run
      MEDIUM → yellow-highlighted run
      LOW    → red-highlighted, strikethrough run

    Parameters
    ----------
    consensus_md_path : Path
        Path to the consensus ``.md`` file.
    stem : str
        Base filename stem.

    Returns
    -------
    Path
        Path to the written ``.docx`` file.
    """
    try:
        from docx import Document  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import RGBColor  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install with: pip install python-docx"
        ) from exc

    md_text = consensus_md_path.read_text(encoding="utf-8")
    doc = Document()

    # Document title style
    doc.core_properties.title = "Chorus Consensus Transcript"
    doc.core_properties.author = "Chorus Engine"

    def _set_highlight(run, colour_name: str) -> None:
        """Apply a Word highlight colour to a run via direct XML manipulation."""
        rpr = run._r.get_or_add_rPr()
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), colour_name)
        rpr.append(hl)

    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            p.add_run(_strip_md_markup(stripped[2:]))
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 60)
        else:
            # Body paragraph — parse inline confidence markers
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]

            # Tokenise line into segments: plain | ==medium== | **~~low~~**[...]
            token_re = re.compile(r"(==.+?==|\*\*~~.+?~~\*\*\[.*?\]|[^=*]+)")
            for match in token_re.finditer(stripped):
                token = match.group(0)
                run = para.add_run()

                if token.startswith("==") and token.endswith("=="):
                    # MEDIUM — yellow highlight
                    word = token[2:-2]
                    run.text = word + " "
                    _set_highlight(run, "yellow")

                elif token.startswith("**~~"):
                    # LOW — red highlight + strikethrough
                    inner = re.sub(r"\*\*~~(.+?)~~\*\*\[.*?\]", r"\1", token)
                    run.text = inner + " "
                    run.font.strike = True
                    run.font.color.rgb = RGBColor(0x72, 0x1C, 0x24)
                    _set_highlight(run, "red")

                else:
                    run.text = token

    target_dir = output_dir or CONSENSUS_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / f"{stem}_consensus.docx"
    doc.save(str(out_path))
    logger.info("DOCX export written → %s", out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# SRT export
# ─────────────────────────────────────────────────────────────────────────────


def _extract_word_timestamps(whisper_result: dict[str, Any]) -> list[dict[str, Any]]:
    """Extract word-level timestamps from Whisper result if available."""
    words: list[dict[str, Any]] = []
    for seg in whisper_result.get("segments", []):
        seg_words = seg.get("words", [])
        for w in seg_words:
            if "start" in w and "end" in w and "word" in w:
                words.append(w)
    return words


def export_srt(
    whisper_result: dict[str, Any],
    stem: str,
    word_level: bool = True,
    output_dir: Path | None = None,
) -> Path:
    """
    Export Whisper data as a SubRip (.srt) subtitle file.

    When *word_level* is True and word timestamps are available, each subtitle
    cue contains a single word with precise timing. Otherwise, falls back to
    segment-level subtitles.

    Parameters
    ----------
    whisper_result : dict
        Whisper result dict containing ``"segments"`` (and optionally
        word-level timestamps within each segment).
    stem : str
        Base filename stem.
    word_level : bool
        If True, use per-word timestamps when available.

    Returns
    -------
    Path
        Path to the written ``.srt`` file.
    """
    lines: list[str] = []

    word_ts = _extract_word_timestamps(whisper_result) if word_level else []

    if word_ts:
        # Word-level SRT: group words into short cues (max 6 words per cue)
        max_words_per_cue = 6
        for cue_idx in range(0, len(word_ts), max_words_per_cue):
            chunk = word_ts[cue_idx : cue_idx + max_words_per_cue]
            start_ts = _seconds_to_srt_ts(chunk[0]["start"])
            end_ts = _seconds_to_srt_ts(chunk[-1]["end"])
            text = " ".join(w["word"].strip() for w in chunk)
            idx = cue_idx // max_words_per_cue + 1
            lines += [str(idx), f"{start_ts} --> {end_ts}", text, ""]
    else:
        # Segment-level fallback
        for idx, seg in enumerate(whisper_result.get("segments", []), start=1):
            start_ts = _seconds_to_srt_ts(seg["start"])
            end_ts = _seconds_to_srt_ts(seg["end"])
            text = _strip_md_markup(seg["text"].strip())
            lines += [str(idx), f"{start_ts} --> {end_ts}", text, ""]

    target_dir = output_dir or CONSENSUS_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / f"{stem}_consensus.srt"
    out_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info("SRT export written → %s", out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# VTT export
# ─────────────────────────────────────────────────────────────────────────────


def export_vtt(
    whisper_result: dict[str, Any],
    stem: str,
    word_level: bool = True,
    output_dir: Path | None = None,
) -> Path:
    """
    Export Whisper data as a WebVTT (.vtt) subtitle file.

    When *word_level* is True and word timestamps are available, each cue
    contains a small group of words with precise timing.

    Parameters
    ----------
    whisper_result : dict
        Whisper result dict.
    stem : str
        Base filename stem.
    word_level : bool
        If True, use per-word timestamps when available.

    Returns
    -------
    Path
        Path to the written ``.vtt`` file.
    """
    lines: list[str] = ["WEBVTT", ""]

    word_ts = _extract_word_timestamps(whisper_result) if word_level else []

    if word_ts:
        max_words_per_cue = 6
        for cue_idx in range(0, len(word_ts), max_words_per_cue):
            chunk = word_ts[cue_idx : cue_idx + max_words_per_cue]
            start_ts = _seconds_to_vtt_ts(chunk[0]["start"])
            end_ts = _seconds_to_vtt_ts(chunk[-1]["end"])
            text = " ".join(w["word"].strip() for w in chunk)
            idx = cue_idx // max_words_per_cue + 1
            lines += [f"{idx}", f"{start_ts} --> {end_ts}", text, ""]
    else:
        for idx, seg in enumerate(whisper_result.get("segments", []), start=1):
            start_ts = _seconds_to_vtt_ts(seg["start"])
            end_ts = _seconds_to_vtt_ts(seg["end"])
            text = _strip_md_markup(seg["text"].strip())
            lines += [f"{idx}", f"{start_ts} --> {end_ts}", text, ""]

    target_dir = output_dir or CONSENSUS_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / f"{stem}_consensus.vtt"
    out_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info("VTT export written → %s", out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# Convenience: export all formats
# ─────────────────────────────────────────────────────────────────────────────


def export_all(
    consensus_md_path: Path,
    whisper_result: dict[str, Any],
    stem: str,
    formats: list[str] | None = None,
    output_dir: Path | None = None,
) -> dict[str, Path | None]:
    """
    Export the consensus document in all requested formats.

    Parameters
    ----------
    consensus_md_path : Path
        Path to the consensus ``.md`` file.
    whisper_result : dict
        Whisper result dict (used for SRT/VTT timestamps).
    stem : str
        Base filename stem.
    formats : list[str], optional
        Subset of ``["pdf", "docx", "srt", "vtt"]``.
        Defaults to all four.

    Returns
    -------
    dict[str, Path | None]
        Mapping of format name → output path (or None if export failed).
    """
    formats = formats or ["pdf", "docx", "srt", "vtt"]
    results: dict[str, Path | None] = {}

    dispatch = {
        "pdf": lambda: export_pdf(consensus_md_path, stem, output_dir=output_dir),
        "docx": lambda: export_docx(consensus_md_path, stem, output_dir=output_dir),
        "srt": lambda: export_srt(whisper_result, stem, output_dir=output_dir),
        "vtt": lambda: export_vtt(whisper_result, stem, output_dir=output_dir),
    }

    for fmt in formats:
        if fmt not in dispatch:
            logger.warning("Unknown export format '%s' — skipping.", fmt)
            results[fmt] = None
            continue
        try:
            results[fmt] = dispatch[fmt]()
        except Exception as exc:
            logger.error("Export failed for format '%s': %s", fmt, exc)
            results[fmt] = None

    return results


# ─────────────────────────────────────────────────────────────────────────────
# Plain-text "most likely" transcript
# ─────────────────────────────────────────────────────────────────────────────


def export_plain_text(
    consensus_md_path: Path,
    stem: str,
    output_dir: Path | None = None,
    include_low: bool = True,
) -> Path:
    """
    Strip Chorus confidence markup from the consensus document and write a
    clean plain-text transcript.

    The body section (between ``## Consensus Transcript`` and the next ``---``)
    is extracted, then confidence decorators are removed:

    - MEDIUM (``==word==``) → ``word``
    - LOW (``**~~word~~**[^…]``) → ``[word?]`` when *include_low* is True,
      or omitted entirely when False
    - HIGH words are already plain text — no change needed

    Parameters
    ----------
    consensus_md_path : Path
        Path to the consensus ``.md`` file.
    stem : str
        Base filename stem.
    include_low : bool
        When True, LOW-confidence words appear as ``[word?]``.
        When False, they are omitted entirely.

    Returns
    -------
    Path
        Path to the written ``.txt`` file.
    """
    text = consensus_md_path.read_text(encoding="utf-8")

    # Extract the transcript body between the heading and the next divider
    match = re.search(
        r"## Consensus Transcript\s*\n\n(.*?)(?=^---)",
        text,
        re.DOTALL | re.MULTILINE,
    )
    body = match.group(1).strip() if match else text

    # MEDIUM: ==word== → word
    body = re.sub(r"==([^=]+)==", r"\1", body)

    # LOW: **~~word~~**[^…] → [word?] or omit
    if include_low:
        body = re.sub(r"\*\*~~([^~]+)~~\*\*\[.*?\]", r"[\1?]", body)
    else:
        body = re.sub(r"\*\*~~[^~]+~~\*\*\[.*?\]", "", body)

    # Collapse multiple spaces left by omissions
    body = re.sub(r" {2,}", " ", body).strip()

    filename = (
        f"{stem}_most_likely.txt" if include_low else f"{stem}_most_likely_clean.txt"
    )
    target_dir = output_dir or CONSENSUS_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / filename
    out_path.write_text(body, encoding="utf-8")
    logger.info("Plain-text export written → %s", out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# Download-all zip bundle
# ─────────────────────────────────────────────────────────────────────────────


def export_zip(
    consensus_md_path: Path,
    whisper_result: dict[str, Any],
    stem: str,
    include_formats: list[str] | None = None,
) -> bytes:
    """
    Bundle all outputs for a recording into an in-memory zip archive.

    Always includes:
    - ``{stem}_consensus.md`` — annotated consensus document
    - ``{stem}_most_likely.txt`` — plain transcript with LOW words in ``[brackets]``
    - ``{stem}_most_likely_clean.txt`` — plain transcript with LOW words omitted

    Optionally includes any formats from ``include_formats``
    (``"pdf"``, ``"docx"``, ``"srt"``, ``"vtt"``).

    Parameters
    ----------
    consensus_md_path : Path
        Path to the consensus ``.md`` file.
    whisper_result : dict
        Whisper result dict (used for SRT/VTT timestamps if requested).
    stem : str
        Base filename stem.
    include_formats : list[str], optional
        Additional export formats to include.

    Returns
    -------
    bytes
        Raw zip archive bytes, ready for ``st.download_button``.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # Consensus markdown — always included
        if consensus_md_path.exists():
            zf.write(consensus_md_path, consensus_md_path.name)

        # Speaker names sidecar — included if it exists
        speaker_names_path = CONSENSUS_DIR / f"{stem}_speakers.json"
        if speaker_names_path.exists():
            zf.write(speaker_names_path, speaker_names_path.name)

        # AI context pack — included if it exists
        ai_context_path = CONSENSUS_DIR / f"{stem}_ai_context.md"
        if ai_context_path.exists():
            zf.write(ai_context_path, ai_context_path.name)

        # Diarised transcript — included if it exists
        diarised_path = CONSENSUS_DIR / f"{stem}_diarised.md"
        if diarised_path.exists():
            zf.write(diarised_path, diarised_path.name)

        # Additional format exports
        if include_formats:
            for _, path in export_all(
                consensus_md_path, whisper_result, stem, include_formats
            ).items():
                if path and path.exists():
                    zf.write(path, path.name)

        # Both plain-text variants — always included
        for include_low in (True, False):
            plain = export_plain_text(consensus_md_path, stem, include_low=include_low)
            if plain.exists():
                zf.write(plain, plain.name)

    buf.seek(0)
    return buf.read()


# ───────────────────────────────────────────────────────────────────────────
# JSON transcript bundle
# ───────────────────────────────────────────────────────────────────────────


def export_transcript_bundle(
    transcripts: dict[str, dict[str, Any]],
    votes: list[Any],
    stem: str,
    *,
    source_filename: str | None = None,
    output_dir: Path | None = None,
) -> Path:
    """
    Write a structured JSON transcript bundle for programmatic consumption.

    The bundle contains:
    - ``meta``           — stem, chorus version, timestamp, source filename
    - ``variants``       — all variant transcript dicts (text, language, model, device)
    - ``consensus``      — word-vote sequence with tier/confidence per word
    - ``statistics``     — HIGH/MEDIUM/LOW word counts and percentages

    Parameters
    ----------
    transcripts : dict[str, dict]
        Mapping of variant key → Whisper result dict.
    votes : list[WordVote]
        Ordered consensus vote sequence from the alignment stage.
    stem : str
        Base filename stem.
    source_filename : str, optional
        Original source filename (with extension) for traceability.
    output_dir : Path, optional
        Directory to write the bundle.  Defaults to ``CONSENSUS_DIR``.

    Returns
    -------
    Path
        Path to the written ``{stem}_bundle.json`` file.
    """
    import json
    from datetime import UTC, datetime

    out_dir = output_dir if output_dir is not None else CONSENSUS_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    high = sum(1 for v in votes if v.tier == "HIGH")
    medium = sum(1 for v in votes if v.tier == "MEDIUM")
    low = sum(1 for v in votes if v.tier == "LOW")
    total_words = max(len(votes), 1)

    bundle: dict[str, Any] = {
        "meta": {
            "stem": stem,
            "source_filename": source_filename,
            "generated_at": datetime.now(UTC).isoformat(),
        },
        "variants": {
            key: {
                "text": result.get("text", ""),
                "language": result.get("language", ""),
                "model": result.get("model", ""),
                "device": result.get("device", ""),
            }
            for key, result in transcripts.items()
        },
        "consensus": [
            {
                "word": v.word,
                "tier": v.tier,
                "confidence": v.confidence,
                "count": v.count,
                "total": v.total,
                "variants": v.variants,
            }
            for v in votes
        ],
        "statistics": {
            "total_words": total_words,
            "high": high,
            "medium": medium,
            "low": low,
            "high_pct": round(high / total_words * 100, 1),
            "medium_pct": round(medium / total_words * 100, 1),
            "low_pct": round(low / total_words * 100, 1),
        },
    }

    out_path = out_dir / f"{stem}_bundle.json"
    with open(out_path, "w", encoding="utf-8") as fh:
        json.dump(bundle, fh, ensure_ascii=False, indent=2)

    logger.info("Transcript bundle written \u2192 %s", out_path)
    return out_path
